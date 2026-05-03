"""
报告格式转换器 —— 将 Markdown 报告转换为 DOCX 和 PDF
"""
import os
import re

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownToDocx:
    """将 Markdown 实验报告转换为排版精良的 DOCX 文件"""

    def __init__(self, markdown_text):
        self.md = markdown_text
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.paragraph_format.line_spacing = 1.5

    def _set_cell_border(self, cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                element = OxmlElement(f'w:{edge}')
                for attr in ['sz', 'val', 'color', 'space']:
                    if attr in edge_data:
                        element.set(qn(f'w:{attr}'), str(edge_data[attr]))
                tcBorders.append(element)
        tcPr.append(tcBorders)

    def _add_borders_to_table(self, table):
        border_data = {
            'top': {'sz': '4', 'val': 'single', 'color': '000000'},
            'start': {'sz': '4', 'val': 'single', 'color': '000000'},
            'bottom': {'sz': '4', 'val': 'single', 'color': '000000'},
            'end': {'sz': '4', 'val': 'single', 'color': '000000'},
            'insideH': {'sz': '4', 'val': 'single', 'color': '000000'},
            'insideV': {'sz': '4', 'val': 'single', 'color': '000000'},
        }
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell, **border_data)

    def _add_run(self, paragraph, text, bold=False, size=None, font_name=None, color=None):
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
        if font_name:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return run

    def convert(self):
        lines = self.md.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            if not line.strip():
                i += 1
                continue

            if line.startswith('# ') and not line.startswith('## '):
                i = self._handle_h1(lines, i)
            elif line.startswith('## ') and not line.startswith('### '):
                i = self._handle_h2(lines, i)
            elif line.startswith('### '):
                i = self._handle_h3(lines, i)
            elif line.startswith('|'):
                i = self._handle_table(lines, i)
            elif re.match(r'^\d+\.\s', line.strip()):
                i = self._handle_numbered_list(lines, i)
            elif line.strip().startswith('- '):
                i = self._handle_bullet_list(lines, i)
            elif line.strip().startswith('> '):
                i = self._handle_blockquote(lines, i)
            elif line.strip().startswith('[') and '在此插入' in line:
                i = self._handle_placeholder(lines, i)
            elif line.strip().startswith('---'):
                i += 1
            else:
                i = self._handle_paragraph(lines, i)

        return self.doc

    def _handle_h1(self, lines, i):
        text = lines[i][2:].strip()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(p, text, bold=True, size=18, font_name='黑体')
        p.paragraph_format.space_after = Pt(12)
        return i + 1

    def _handle_h2(self, lines, i):
        text = lines[i][3:].strip()
        p = self.doc.add_paragraph()
        self._add_run(p, text, bold=True, size=15, font_name='黑体')
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return i + 1

    def _handle_h3(self, lines, i):
        text = lines[i][4:].strip()
        p = self.doc.add_paragraph()
        self._add_run(p, text, bold=True, size=13, font_name='黑体')
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        return i + 1

    def _handle_paragraph(self, lines, i):
        text = lines[i].strip()

        if text.startswith('**') and text.endswith('**'):
            text = text[2:-2]
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_run(p, text, bold=True, size=12, font_name='黑体')
            return i + 1

        if not text:
            return i + 1

        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)

        self._parse_inline_formatting(p, text)
        return i + 1

    def _parse_inline_formatting(self, paragraph, text):
        segments = re.split(r'(\*\*.*?\*\*)', text)
        for seg in segments:
            if seg.startswith('**') and seg.endswith('**'):
                self._add_run(paragraph, seg[2:-2], bold=True)
            else:
                self._add_run(paragraph, seg)

    def _handle_numbered_list(self, lines, i):
        while i < len(lines):
            line = lines[i].strip()
            m = re.match(r'^(\d+)\.\s(.*)', line)
            if not m:
                break
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            self._add_run(p, f"{m.group(1)}. {m.group(2)}")
            i += 1
        return i

    def _handle_bullet_list(self, lines, i):
        while i < len(lines):
            line = lines[i].strip()
            if not line.startswith('- '):
                break
            text = line[2:]
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            self._add_run(p, f"• {text}")
            i += 1
        return i

    def _handle_blockquote(self, lines, i):
        while i < len(lines):
            line = lines[i].strip()
            if not line.startswith('> '):
                break
            text = line[2:]
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            self._add_run(p, text, size=10.5, color=(100, 100, 100))
            i += 1
        return i

    def _handle_placeholder(self, lines, i):
        text = lines[i].strip()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(p, text, bold=True, size=12, color=(180, 0, 0))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        return i + 1

    def _handle_table(self, lines, i):
        header_line = lines[i].strip()
        header_cells = [c.strip() for c in header_line.split('|')[1:-1]]

        i += 1
        if i < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i]):
            i += 1

        data_rows = []
        while i < len(lines):
            line = lines[i].strip()
            if not line.startswith('|'):
                break
            row_cells = [c.strip() for c in line.split('|')[1:-1]]
            data_rows.append(row_cells)
            i += 1

        num_rows = 1 + len(data_rows)
        num_cols = len(header_cells)
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_borders_to_table(table)

        for j, cell_text in enumerate(header_cells):
            cell = table.rows[0].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_run(p, cell_text, bold=True, size=10.5, font_name='黑体')
            self._set_cell_shading(cell, 'D9E2F3')

        for r, row_data in enumerate(data_rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[r + 1].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._add_run(p, cell_text, size=10.5)

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        return i

    def _set_cell_shading(self, cell, color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        shading_elm.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading_elm)


def markdown_to_docx(markdown_text, output_path):
    converter = MarkdownToDocx(markdown_text)
    doc = converter.convert()
    doc.save(output_path)
    return output_path


def docx_to_pdf(docx_path, pdf_path=None):
    from docx2pdf import convert
    if pdf_path is None:
        pdf_path = docx_path.replace('.docx', '.pdf')
    convert(docx_path, pdf_path)
    return pdf_path


def generate_formats(markdown_text, base_output_dir, base_filename):
    os.makedirs(base_output_dir, exist_ok=True)

    docx_path = os.path.join(base_output_dir, f"{base_filename}.docx")
    markdown_to_docx(markdown_text, docx_path)
    print(f"DOCX 已生成: {docx_path}")

    pdf_path = os.path.join(base_output_dir, f"{base_filename}.pdf")
    try:
        docx_to_pdf(docx_path, pdf_path)
        print(f"PDF 已生成: {pdf_path}")
    except Exception as e:
        print(f"PDF 生成失败 (可能需要安装 Microsoft Word): {e}")
        print(f"替代方案: 请手动用 Word 打开 DOCX 文件并另存为 PDF。")

    return docx_path, pdf_path
