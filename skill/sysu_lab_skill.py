"""
SYSU_Lab_Report_Gen 排版技能
中山大学电子与信息工程学院高频电路实验报告排版铁律

本模块提供排版常量和工具函数，所有实验报告生成器必须无条件调用。
"""
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class SYSULabReportSkill:
    """中山大学高频电路实验报告排版铁律"""

    # === 字体 ===
    FONT_BODY_CN = '宋体'
    FONT_TITLE_CN = '黑体'
    FONT_WESTERN = 'Times New Roman'

    # === 字号 ===
    SIZE_MAIN_TITLE = Pt(26)       # 一号，封面大标题
    SIZE_EXP_TITLE = Pt(15)       # 小三号，实验小标题（单行）
    SIZE_SECTION = Pt(15)         # 小三，章节标题
    SIZE_SUB_SECTION = Pt(14)     # 四号，子标题
    SIZE_BODY = Pt(12)            # 小四，正文
    SIZE_CAPTION = Pt(10.5)       # 五号，图注
    SIZE_TABLE = Pt(9)            # 小五，表格内文字
    SIZE_INFO = Pt(10.5)       # 五号，基础信息

    # === 颜色 ===
    COLOR_CAPTION = RGBColor(0x55, 0x55, 0x55)
    COLOR_PLACEHOLDER = RGBColor(0x99, 0x99, 0x99)
    COLOR_IMAGE_BOX_BORDER = '999999'
    COLOR_TABLE_HEADER_BG = 'D9E2F3'
    COLOR_HORIZONTAL_RULE = '000000'

    # === 尺寸 ===
    TABLE_ROW_HEIGHT_MIN = Cm(0.8)   # 数据表格最小行高，防手写拥挤
    IMAGE_BOX_HEIGHT = Cm(6.0)       # 图片占位框高度
    IMAGE_BOX_BORDER_WIDTH = '6'     # 图片框边框宽度，单位 1/8 pt
    FIRST_LINE_INDENT = Cm(0.74)     # 首行缩进 2 字符

    # === 页面设置 ===
    PAGE_WIDTH = Cm(21.0)
    PAGE_HEIGHT = Cm(29.7)
    MARGIN_TOP = Cm(2.54)
    MARGIN_BOTTOM = Cm(2.54)
    MARGIN_LEFT = Cm(3.18)
    MARGIN_RIGHT = Cm(3.18)

    # === 页眉页脚 ===
    HEADER_TEXT = '中山大学电子与信息工程学院高频电路实验报告'
    HEADER_FONT_SIZE = Pt(10.5)
    FOOTER_FONT_SIZE = Pt(9)


# 将类作为单例导出，方便各处引用
SKILL = SYSULabReportSkill()


def set_run_font(run, cn_font=SKILL.FONT_BODY_CN, en_font=SKILL.FONT_WESTERN,
                 size=SKILL.SIZE_BODY, bold=False, color=None):
    """精确设置 run 的中英文字体"""
    run.font.name = en_font
    run.font.size = size
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:cs'), en_font)
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=0,
                          first_line_indent=None, alignment=None):
    """设置段落间距和对齐"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    if alignment is not None:
        paragraph.alignment = alignment


def add_centered_run_paragraph(doc, text, cn_font=SKILL.FONT_BODY_CN,
                               size=SKILL.SIZE_BODY, bold=False, color=None,
                               space_before=0, space_after=0):
    """添加居中段落"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.5, space_before=space_before,
                          space_after=space_after, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size, bold=bold, color=color)
    return p


def add_left_paragraph(doc, text, cn_font=SKILL.FONT_BODY_CN,
                       size=SKILL.SIZE_BODY, bold=False, color=None,
                       space_before=0, space_after=0, indent=True):
    """添加左对齐正文段落，带首行缩进"""
    p = doc.add_paragraph()
    fi = SKILL.FIRST_LINE_INDENT if indent else None
    set_paragraph_spacing(p, line_spacing=1.5, space_before=space_before,
                          space_after=space_after, first_line_indent=fi,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size, bold=bold, color=color)
    return p


def add_section_title(doc, text, size=SKILL.SIZE_SECTION, space_before=14, space_after=8):
    """添加节标题：黑体加粗"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.5, space_before=space_before,
                          space_after=space_after, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_run_font(run, cn_font=SKILL.FONT_TITLE_CN, size=size, bold=True)
    return p


def add_sub_section_title(doc, text, size=SKILL.SIZE_SUB_SECTION, space_before=10, space_after=6):
    """添加子节标题：黑体加粗"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.5, space_before=space_before,
                          space_after=space_after, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_run_font(run, cn_font=SKILL.FONT_TITLE_CN, size=size, bold=True)
    return p


def add_horizontal_rule(doc, space_before=4, space_after=6):
    """绘制贯穿页面的黑色实线水平分割线"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.0, space_before=space_before,
                          space_after=space_after, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), SKILL.COLOR_HORIZONTAL_RULE)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr, val in edge_data.items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_row_height(row, height):
    """强制设置行高"""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height.emu / 635)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def get_border_spec():
    """标准表格边框规格"""
    return {
        'top': {'sz': '4', 'val': 'single', 'color': '000000'},
        'start': {'sz': '4', 'val': 'single', 'color': '000000'},
        'bottom': {'sz': '4', 'val': 'single', 'color': '000000'},
        'end': {'sz': '4', 'val': 'single', 'color': '000000'},
        'insideH': {'sz': '4', 'val': 'single', 'color': '000000'},
        'insideV': {'sz': '4', 'val': 'single', 'color': '000000'},
    }


def add_data_table(doc, headers, rows, col_width=None):
    """
    添加标准数据表格。
    - 表头：蓝色背景，黑体小五加粗
    - 数据行：强制最小行高 0.8cm
    - 全边框
    """
    num_cols = len(headers)
    num_rows = 1 + len(rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.autofit = True

    border = get_border_spec()
    for row_obj in table.rows:
        for cell in row_obj.cells:
            set_cell_border(cell, **border)

    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        _set_cell_text(cell, hdr, cn_font=SKILL.FONT_TITLE_CN,
                       size=SKILL.SIZE_TABLE, bold=True, align='center')
        set_cell_shading(cell, SKILL.COLOR_TABLE_HEADER_BG)

    for r, row_data in enumerate(rows):
        row_obj = table.rows[r + 1]
        set_row_height(row_obj, SKILL.TABLE_ROW_HEIGHT_MIN)
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row_obj.cells[j]
                text = cell_text if cell_text else '——'
                _set_cell_text(cell, text, size=SKILL.SIZE_TABLE, align='center')

    if col_width:
        for row_obj in table.rows:
            for j, w in enumerate(col_width):
                if j < num_cols:
                    row_obj.cells[j].width = Cm(w)

    return table


def _set_cell_text(cell, text, cn_font=SKILL.FONT_BODY_CN,
                   size=SKILL.SIZE_TABLE, bold=False, align='center'):
    """设置单元格文字"""
    cell.text = ''
    p = cell.paragraphs[0]
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, line_spacing=1.0, space_before=1, space_after=1)
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size, bold=bold)


def add_image_placeholder_box(doc, fig_number, description):
    """
    图片占位符 —— 定格表格
    1 行 1 列，宽度 100% 页宽，高度固定 6cm
    内部居中灰色提示文字
    """
    table = doc.add_table(rows=1, cols=1)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.autofit = True

    border = {
        'top': {'sz': SKILL.IMAGE_BOX_BORDER_WIDTH, 'val': 'single',
                'color': SKILL.COLOR_IMAGE_BOX_BORDER},
        'start': {'sz': SKILL.IMAGE_BOX_BORDER_WIDTH, 'val': 'single',
                  'color': SKILL.COLOR_IMAGE_BOX_BORDER},
        'bottom': {'sz': SKILL.IMAGE_BOX_BORDER_WIDTH, 'val': 'single',
                   'color': SKILL.COLOR_IMAGE_BOX_BORDER},
        'end': {'sz': SKILL.IMAGE_BOX_BORDER_WIDTH, 'val': 'single',
                'color': SKILL.COLOR_IMAGE_BOX_BORDER},
    }
    cell = table.rows[0].cells[0]
    set_cell_border(cell, **border)
    set_row_height(table.rows[0], SKILL.IMAGE_BOX_HEIGHT)

    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_cell_vertical_alignment(cell, 'center')

    set_paragraph_spacing(p, line_spacing=1.2, space_before=0, space_after=0)
    run = p.add_run(f'[ 请在此框内粘贴  图 {fig_number} - {description}  截图 ]')
    set_run_font(run, size=Pt(10.5), color=SKILL.COLOR_PLACEHOLDER)

    return table


def set_cell_vertical_alignment(cell, align='center'):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def setup_page(doc):
    """设置 A4 页面和页边距"""
    section = doc.sections[0]
    section.page_width = SKILL.PAGE_WIDTH
    section.page_height = SKILL.PAGE_HEIGHT
    section.top_margin = SKILL.MARGIN_TOP
    section.bottom_margin = SKILL.MARGIN_BOTTOM
    section.left_margin = SKILL.MARGIN_LEFT
    section.right_margin = SKILL.MARGIN_RIGHT


def setup_header(section, experiment_title=''):
    """设置页眉：小五号宋体居中 + 底部边框线"""
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(hp, line_spacing=1.0, space_before=0, space_after=0)

    header_text = f'高频电路实验报告  {experiment_title}' if experiment_title else SKILL.HEADER_TEXT
    run = hp.add_run(header_text)
    set_run_font(run, size=Pt(9))

    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def setup_footer(section):
    """设置页脚：居中页码 - N -"""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(fp, line_spacing=1.0, space_before=0, space_after=0)

    run1 = fp.add_run('- ')
    set_run_font(run1, size=SKILL.FOOTER_FONT_SIZE)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run2 = fp.add_run()
    run2._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run3 = fp.add_run()
    run3._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run4 = fp.add_run()
    run4._r.append(fldChar2)

    run5 = fp.add_run(' -')
    set_run_font(run5, size=SKILL.FOOTER_FONT_SIZE)


def add_figure_caption(doc, text, fig_number):
    """添加图注：五号宋体居中深灰色"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.5, space_before=Pt(2),
                          space_after=Pt(8), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(f'图 {fig_number}  {text}')
    set_run_font(run, size=SKILL.SIZE_CAPTION, color=SKILL.COLOR_CAPTION)
    return p


def add_blank_line(doc, size=Pt(6)):
    """添加空白分隔行"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.0, space_before=0, space_after=0)
    run = p.add_run('')
    set_run_font(run, size=size)
    return p


def add_info_table(doc):
    """
    基础信息区 —— 3行×2列 无边框隐形表格，绝对对齐
    第1行: 学院：...    | 专业：...（55:45列宽，宽左窄右黄金比例）
    第2行: 实验人：... | 参加人：...
    第3行: 日期：...    | (空)
    五号宋体，左列55%宽→中心线右移，右列自然靠右，视觉对称。
    """
    table = doc.add_table(rows=3, cols=2)
    table.autofit = True

    from docx.enum.table import WD_TABLE_ALIGNMENT
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    none_border = {
        'top': {'sz': '0', 'val': 'nil', 'color': 'auto'},
        'start': {'sz': '0', 'val': 'nil', 'color': 'auto'},
        'bottom': {'sz': '0', 'val': 'nil', 'color': 'auto'},
        'end': {'sz': '0', 'val': 'nil', 'color': 'auto'},
    }
    for row_obj in table.rows:
        for cell in row_obj.cells:
            set_cell_border(cell, **none_border)

    page_text_width = SKILL.PAGE_WIDTH - SKILL.MARGIN_LEFT - SKILL.MARGIN_RIGHT
    left_col_width = page_text_width * 0.55
    right_col_width = page_text_width * 0.45
    for row_obj in table.rows:
        row_obj.cells[0].width = left_col_width
        row_obj.cells[1].width = right_col_width

    data = [
        ('学院：电子与信息工程学院', '专业：电子信息科学与技术'),
        ('实验人：__________________', '参加人：__________________'),
        ('日期：__________________', ''),
    ]

    for r, (left_text, right_text) in enumerate(data):
        _set_cell_text(table.rows[r].cells[0], left_text,
                       size=SKILL.SIZE_INFO, align='left')
        _set_cell_text(table.rows[r].cells[1], right_text,
                       size=SKILL.SIZE_INFO, align='left')

    set_paragraph_spacing(doc.add_paragraph(), line_spacing=1.0, space_before=0, space_after=0)

    return table


def init_document(experiment_title=''):
    """初始化文档：创建 Document，设全局样式、页面、页眉页脚"""
    from docx import Document
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = SKILL.FONT_WESTERN
    style.font.size = SKILL.SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), SKILL.FONT_BODY_CN)
    style.paragraph_format.line_spacing = 1.5

    setup_page(doc)
    setup_header(doc.sections[0], experiment_title=experiment_title)
    setup_footer(doc.sections[0])

    return doc
